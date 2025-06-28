from pathlib import Path
import docx
import pandas as pd
from tqdm import tqdm
import time
import yaml
from fnmatch import fnmatch
import os
import logging
import glob

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load config
with open("config.yaml", "r") as f:
    config = yaml.safe_load(f)
file_patterns = config.get("file_patterns", None)

def extract_docx_tables(doc):
    """Extract tables from a DOCX document and format them as text."""
    table_texts = []
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            # Get cell text and handle empty cells
            row_text = [cell.text.strip() if cell.text.strip() else "" for cell in row.cells]
            table_text.append(" | ".join(row_text))
        table_texts.append("\n".join(table_text))
    return "\n\n".join(table_texts)

def parse_files_and_web_page(directory, web_page_file):
    participants_rows = []
    participants_headers = None
    definitions_blocks = []
    print("\nStarting data processing for LlamaIndex...")

    # Add web page content
    if web_page_file:
        try:
            print("\nProcessing web page content...")
            with open(web_page_file, "r", encoding="utf-8") as f:
                page_content = f.read()
            print("✓ Web page content processed successfully")
        except Exception as e:
            print(f"Error reading web page file {web_page_file}: {e}")

    # Get list of files to process
    all_files = list(Path(directory).glob("*"))
    if file_patterns:
        files = [f for f in all_files if any(fnmatch(f.name, pat) for pat in file_patterns)]
    else:
        files = all_files
    print(f"\nFound {len(files)} files to process")

    for file_path in tqdm(files, desc="Processing files", unit="file"):
        if file_path.suffix.lower() == ".csv":
            try:
                encodings = ["utf-8", "latin1", "cp1252", "iso-8859-1"]
                df = None
                for encoding in encodings:
                    try:
                        df = pd.read_csv(
                            file_path,
                            encoding=encoding,
                            low_memory=False,
                            dtype=str,
                        )
                        break
                    except UnicodeDecodeError:
                        continue
                if df is not None:
                    if participants_headers is None:
                        participants_headers = df.columns.tolist()
                    for _, row in df.iterrows():
                        participants_rows.append([str(val) for val in row])
                else:
                    print(f"Could not read CSV {file_path} with any of the attempted encodings")
            except Exception as e:
                print(f"Error reading CSV {file_path}: {e}")
        elif file_path.suffix.lower() == ".docx":
            try:
                doc = docx.Document(file_path)
                for table in doc.tables:
                    table_lines = []
                    for row in table.rows:
                        row_text = [cell.text.strip() if cell.text.strip() else "" for cell in row.cells]
                        table_lines.append(" | ".join(row_text))
                    definitions_blocks.append("\n".join(table_lines))
            except Exception as e:
                print(f"Error reading DOCX {file_path}: {e}")

    # Output participants_data.csv
    output_dir = Path("data/processed")
    output_dir.mkdir(parents=True, exist_ok=True)
    participants_csv = output_dir / "participants_data.csv"
    if participants_headers is not None:
        with open(participants_csv, "w", encoding="utf-8") as f:
            f.write(",".join(participants_headers) + "\n")
            for row in participants_rows:
                f.write(",".join(row) + "\n")
        print(f"✓ Participants data saved to {participants_csv}")
    else:
        print("No participant data found.")

    # Output definitions.txt
    definitions_txt = output_dir / "definitions.txt"
    if definitions_blocks:
        with open(definitions_txt, "w", encoding="utf-8") as f:
            for block in definitions_blocks:
                f.write(block + "\n\n---\n\n")
        print(f"✓ Definitions saved to {definitions_txt}")
    else:
        print("No definitions found.")

    return str(participants_csv), str(definitions_txt)

def load_config():
    """Load and validate configuration."""
    config_path = Path(__file__).parent.parent / "config.yaml"
    try:
        with open(config_path, "r") as f:
            config = yaml.safe_load(f)
        
        # Validate required config sections
        required_sections = ["file_patterns", "definition_file", "web_content_file"]
        for section in required_sections:
            if section not in config:
                raise ValueError(f"Missing required section '{section}' in config")
        
        return config
    except Exception as e:
        logger.error(f"Error loading config: {e}")
        raise

def get_base_pattern(pattern):
    """Extract base pattern without wildcards."""
    # Remove wildcards and date patterns
    base = pattern.replace("*", "")
    # Remove date pattern if it exists
    base = base.split("_20")[0] if "_20" in base else base
    return base

def process_csv_files(config):
    """Process and combine CSV files based on patterns."""
    logger.info("Processing CSV files")
    state_filter = config.get("state_filter", None)
    srvc_dstrct_filter = config.get("srvc_dstrct_filter", None)
    # Group patterns by their base name
    pattern_groups = {}
    for pattern in config["file_patterns"]:
        if pattern.endswith(".csv") or "*" in pattern:
            base_pattern = get_base_pattern(pattern)
            if base_pattern not in pattern_groups:
                pattern_groups[base_pattern] = []
            pattern_groups[base_pattern].append(pattern)
    # Process each group
    for base_pattern, patterns in pattern_groups.items():
        logger.info(f"Processing group: {base_pattern}")
        all_data = []
        # Find all matching files for these patterns
        matching_files = set()
        for pattern in patterns:
            # Look in both raw and processed directories
            for directory in ["data/raw", "data/processed"]:
                matching_files.update(glob.glob(f"{directory}/{pattern}"))
        if not matching_files:
            logger.warning(f"No files found matching patterns: {patterns}")
            continue
        logger.info(f"Found {len(matching_files)} files: {matching_files}")
        # Process each matching file
        for file_path in tqdm(matching_files, desc=f"Processing {base_pattern} files", unit="file"):
            if not file_path.lower().endswith('.csv'):
                logger.info(f"Skipping non-CSV file: {file_path}")
                continue
            try:
                logger.info(f"Reading {file_path}")
                # Try different encodings
                encodings = ["utf-8", "latin1", "cp1252", "iso-8859-1"]
                df = None
                for encoding in encodings:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, dtype=str)
                        break
                    except UnicodeDecodeError:
                        continue
                if df is None:
                    logger.error(f"Could not read {file_path} with any encoding")
                    continue
                # Filter by StateCd if state_filter is set and column exists
                if state_filter and "StateCd" in df.columns:
                    df = df[df["StateCd"].isin(state_filter)]
                    logger.info(f"Filtered by StateCd: {state_filter}, remaining rows: {len(df)}")
                # Filter by SrvcDstrctNm if srvc_dstrct_filter is set and column exists
                if srvc_dstrct_filter and "SrvcDstrctNm" in df.columns:
                    df = df[df["SrvcDstrctNm"].isin(srvc_dstrct_filter)]
                    logger.info(f"Filtered by SrvcDstrctNm: {srvc_dstrct_filter}, remaining rows: {len(df)}")
                # Add source file information
                df['source_file'] = Path(file_path).name
                # Basic cleaning
                df = df.fillna("")  # Replace NaN with empty string
                df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)  # Strip whitespace
                all_data.append(df)
                logger.info(f"Successfully read {len(df)} rows from {file_path}")
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                continue
        if all_data:
            try:
                # Combine all dataframes
                combined_df = pd.concat(all_data, ignore_index=True)
                # Sort by date if date column exists
                date_columns = [col for col in combined_df.columns if 'date' in col.lower()]
                if date_columns:
                    combined_df = combined_df.sort_values(by=date_columns[0], ascending=False)
                # Save combined CSV
                output_path = f"data/processed/{base_pattern}.csv"
                combined_df.to_csv(output_path, index=False)
                logger.info(f"Combined {len(combined_df)} rows into {output_path}")
            except Exception as e:
                logger.error(f"Error combining data for {base_pattern}: {e}")
        else:
            logger.warning(f"No data to combine for {base_pattern}")

def process_docx_files(input_dir, output_file):
    """Process DOCX files and extract definitions into a single text file."""
    logger.info(f"Processing DOCX files from {input_dir}")
    all_definitions = []
    
    if not Path(input_dir).exists():
        logger.warning(f"DOCX directory {input_dir} does not exist. Creating it.")
        Path(input_dir).mkdir(parents=True, exist_ok=True)
        return
    
    for file in tqdm(list(Path(input_dir).glob("*.docx")), desc="Processing DOCX files", unit="file"):
        logger.info(f"Processing {file.name}")
        try:
            doc = docx.Document(file)
            
            # Extract tables and their content
            for table in doc.tables:
                definition = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):  # Only add non-empty rows
                        definition.append(" | ".join(row_text))
                
                if definition:  # Only add non-empty tables
                    all_definitions.append("\n".join(definition))
        except Exception as e:
            logger.error(f"Error processing {file.name}: {e}")
            continue
    
    if all_definitions:
        # Write all definitions to file
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n\n---\n\n".join(all_definitions))
        logger.info(f"Definitions written to {output_file}")
    else:
        logger.warning("No definitions found in DOCX files")

def process_web_content(input_file, output_file):
    """Process web content into a clean text file."""
    logger.info(f"Processing web content from {input_file}")
    try:
        if not Path(input_file).exists():
            logger.warning(f"Web content file {input_file} does not exist")
            return
            
        with open(input_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        # Basic cleaning
        content = content.strip()
        
        # Save cleaned content
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"Web content saved to {output_file}")
    except Exception as e:
        logger.error(f"Error processing web content: {e}")

def main():
    try:
        config = load_config()
        
        # Create processed directory if it doesn't exist
        processed_dir = Path("data/processed")
        processed_dir.mkdir(parents=True, exist_ok=True)
        
        # Process CSV files
        process_csv_files(config)
        
        # Process DOCX files
        docx_dir = Path("data/raw")
        definitions_file = processed_dir / "definitions.txt"
        process_docx_files(docx_dir, definitions_file)
        
        # Process web content
        web_input = Path("data/raw/web_page_content.txt")
        web_output = processed_dir / "web_page_content.txt"
        process_web_content(web_input, web_output)
        
        logger.info("All files processed successfully!")
    except Exception as e:
        logger.error(f"Error in main process: {e}")
        raise

if __name__ == "__main__":
    main()
